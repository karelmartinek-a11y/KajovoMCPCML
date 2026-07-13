#!/usr/bin/env python3
"""Create Connect in Catalog v1.4 from the approved v1.2 layout master."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def set_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def all_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def replace_everywhere(document, old: str, new: str) -> int:
    count = 0
    for paragraph in all_paragraphs(document):
        if old in paragraph.text:
            set_text(paragraph, paragraph.text.replace(old, new))
            count += 1
    return count


def replace_exact(document, old: str, new: str) -> None:
    for paragraph in all_paragraphs(document):
        if paragraph.text.strip() == old:
            set_text(paragraph, new)
            return
    raise RuntimeError(f"Paragraph not found: {old}")


def fill_after_heading(document, heading: str, texts: list[str]) -> None:
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() != heading:
            continue
        candidates = document.paragraphs[index + 1:index + 1 + len(texts)]
        if len(candidates) != len(texts) or any(candidate.text.strip() for candidate in candidates):
            raise RuntimeError(f"Expected {len(texts)} empty paragraphs after {heading}")
        for candidate, text in zip(candidates, texts, strict=True):
            set_text(candidate, text)
        return
    raise RuntimeError(f"Heading not found: {heading}")


def build(source: Path, destination: Path) -> None:
    document = Document(source)

    replace_everywhere(document, "Connect in Catalog v1.2", "Connect in Catalog v1.4")
    replace_everywhere(document, "Verze 1.2", "Verze 1.4")
    replace_everywhere(
        document,
        "FAILED, QUARANTINED, CANCELLED nebo expirovaný token jsou nedokončený výsledek; předej správci job ID, correlation ID a blockingError.",
        "Při programmerAction=UPLOAD_REVISION oprav blockingError a neúspěšné gates a nahraj novou revizi. Dokončeno je pouze COMPLETE / ACTIVE; pouze STOP / QUARANTINED předej správci s job ID a correlation ID."
    )
    replace_everywhere(
        document,
        '"toolName": "kcml0042_example_handler"\n  }',
        '"toolName": "kcml0042_example_handler",\n    "programmerAction": {"kind":"WAIT","canUploadRevision":false,\n      "message":"Poll this job until it becomes ACTIVE or requests a new revision."}\n  }'
    )

    replace_exact(
        document,
        "Zadat srozumitelné označení. Systém vytvoří 512bitovou náhodnou hodnotu s prefixem kci_, zobrazí ji právě jednou a do databáze uloží pouze HMAC digest, key ID a šestnáctiznakový fingerprint.",
        "Zadat poznámku k budoucímu serveru. Systém vytvoří 512bitovou náhodnou hodnotu s prefixem kci_, zobrazí ji právě jednou a do databáze uloží pouze HMAC digest, key ID a šestnáctiznakový fingerprint. Poznámka je pouze interní označení a nemění identitu ani konfiguraci serveru."
    )
    replace_exact(
        document,
        "Token předat programátorovi odděleným bezpečným kanálem. Nevkládat jej do repozitáře, ticketu, screenshotu, logu, příkazové historie ani CI proměnné sdílené s pull requesty.",
        "Programátorovi předat tento Connect in Catalog v1.4 a token bezpečným kanálem. Od prvního uploadu programátor sám obsluhuje stav jobu, diagnostiku a opravné revize; správce nemusí ručně vyplňovat registraci, přidělovat HTTPS adresu ani potvrzovat aktivaci. Token nevkládat do repozitáře, ticketu, screenshotu, logu, příkazové historie ani CI proměnné sdílené s pull requesty."
    )
    replace_exact(
        document,
        "V sekci Implementační tokeny sledovat aktuální expiraci, 24hodinový strop, job, KCML kód, PR/CI/deploy/testy a výslednou HTTPS adresu. Token nelze editovat; lze jej revokovat nebo skrýt s ponecháním auditu.",
        "Sekce Implementační tokeny zobrazuje expiraci, job, KCML kód, PR/CI/deploy/testy a výslednou HTTPS adresu pouze pro dohled. U opravitelné chyby jedná programátor přes programátorské API; zásah správce není součástí běžného toku. Token nelze editovat; lze jej revokovat nebo skrýt s ponecháním auditu."
    )

    fill_after_heading(document, "28.2 Polling, opravená revize a zrušení", [
        "Programátor polluje GET téhož jobu. Odpověď vždy obsahuje stav, programmerAction, lockVersion, blokující chybu, jednotlivé gates, correlation ID a bezpečné odkazy na PR/CI. Akce WAIT znamená pokračovat v pollingu; COMPLETE znamená, že server je ACTIVE.",
        "Akce UPLOAD_REVISION znamená opravit blockingError a neúspěšné gates, načíst aktuální ETag/lockVersion a odeslat PUT nové revize se stejným tokenem a jobem, novým Idempotency-Key a If-Match. Identita KCML ani HTTPS adresa se při opravě nemění.",
        "Cyklus GET -> oprava -> PUT -> GET se opakuje bez součinnosti správce až do COMPLETE / ACTIVE. STOP se používá pouze pro CANCELLED nebo neobejitelnou bezpečnostní karanténu; QUARANTINED nesmí programátor sám obnovit ani obejít."
    ])

    replace_exact(
        document,
        "2 hodiny od vystavení. Klientské polling ani upload expiraci neprodlužují.",
        "2 hodiny od vystavení pro první upload. Klientské polling ani upload samy expiraci neprodlužují."
    )
    replace_exact(
        document,
        "Aktivní pronajatý job může nejvýše jednou za 15 minut posunout expiraci na now + 2h, nikdy za issuedAt + 24h.",
        "Po navázání prvního uploadu aktivní serverový job nejvýše jednou za 15 minut posune expiraci na now + 2h, nejdéle do issuedAt + 24h. To zahrnuje čekání na opravitelnou revizi, takže programátor může bez správce dovést běžný opravný cyklus do zeleného výsledku."
    )
    replace_exact(
        document,
        "Po expiraci server zůstane disabled a worker se zastaví. Správce může vystavit navazující token ke stejnému jobu; původní token se revokuje a identita se nemění.",
        "Běžný samoobslužný tok musí skončit do 24hodinového bezpečnostního stropu. Teprve po jeho překročení zůstane server disabled a worker se zastaví; navazující token ke stejné identitě je mimořádná správcovská obnova, nikoli součást standardního opravného cyklu."
    )
    replace_exact(
        document,
        "Token chybí, je chybný, expirovaný, revokovaný nebo nepatří jobu. Nezkoušet obcházet UI; vyžádat resume token.",
        "Token chybí, je chybný, expirovaný, revokovaný nebo nepatří jobu. Opravit předání hlavičky; expirovaný nebo revokovaný token nelze obejít."
    )
    replace_exact(
        document,
        "FAILED / QUARANTINED / CANCELLED",
        "FAILED / QUARANTINED / CANCELLED"
    )
    replace_exact(
        document,
        "Server je vypnut, access a testovací tokeny i egress capability jsou revokovány; QUARANTINED vyžaduje zásah správce.",
        "Server je vypnut. FAILED s programmerAction=UPLOAD_REVISION dovolí opravený zdroj stejného jobu; QUARANTINED je neobejitelná bezpečnostní stopka a CANCELLED je konečný stav."
    )
    replace_exact(
        document,
        "Onboarding worker s GitHub App vytvoří branch integration/kcmlNNNN/<job-id>, zapíše pouze handlers/KCMLNNNN/ a otevře PR. Workflow, root konfigurace, auth a platformní bezpečnostní soubory nesmí PR změnit. GitHub App má read oprávnění pro Actions, aby svázala úspěšný trusted main run ID s provenance.",
        "Onboarding worker s provozovatelem nakonfigurovanou GitHub API autorizací vytvoří branch integration/kcmlNNNN/<job-id>, zapíše pouze handlers/KCMLNNNN/ a otevře PR. Workflow, root konfigurace, auth a platformní bezpečnostní soubory nesmí PR změnit. Autorizace má přístup pouze k potřebným operacím repozitáře a ke čtení Actions, aby svázala úspěšný trusted main run ID s provenance."
    )
    replace_exact(
        document,
        "☐ Pokud job skončil FAILED nebo QUARANTINED, programátor stav nepřepisuje; předá correlation ID a blockingError správci.",
        "☐ Při UPLOAD_REVISION programátor opravil blockingError a neúspěšné gates a opakoval PUT/GET až do ACTIVE. Pouze STOP / QUARANTINED předá správci s job ID a correlation ID."
    )

    core_text = "\n".join(paragraph.text for paragraph in all_paragraphs(document))
    required = [
        "Connect in Catalog v1.4",
        "programmerAction",
        "UPLOAD_REVISION",
        "GET -> oprava -> PUT -> GET",
        "COMPLETE / ACTIVE",
        "QUARANTINED nesmí programátor sám obnovit"
    ]
    missing = [item for item in required if item not in core_text]
    if missing:
        raise RuntimeError(f"Missing required v1.4 content: {missing}")

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("destination", type=Path)
    args = parser.parse_args()
    build(args.source, args.destination)


if __name__ == "__main__":
    main()

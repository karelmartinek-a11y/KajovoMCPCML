from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path("/Volumes/KINGSTON/GITHUB2")
CONNECT_SOURCE = ROOT / "Connect_in_Catalog_KajovoMCPCML_v1.1.docx"
CONNECT_OUTPUT = ROOT / "Connect_in_Catalog_KajovoMCPCML_v1.2.docx"
SSOT_SOURCE = ROOT / "doc/KCML_Spravce_MCP_serveru_SSOT_v1.3_tokeny_min_512_bit.docx"
SSOT_OUTPUT = ROOT / "doc/KCML_Spravce_MCP_serveru_SSOT_v1.4_automaticky_onboarding.docx"

CONNECT_BLUE = "1F4E79"
SSOT_BLUE = "17365D"
LIGHT_BLUE = "E8F1F7"
LIGHT_GRAY = "F3F4F6"
WHITE = "FFFFFF"
TEXT = "1F1F1F"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=110, start=140, bottom=110, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def enable_field_updates(doc):
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths_dxa, indent_dxa=0):
    table.autofit = False
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def style_cell(cell, bold=False, color=TEXT, size=9.6, font="Aptos"):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        for run in paragraph.runs:
            run.font.name = font
            fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            run.font.size = Pt(size)
            run.font.bold = bold
            color_node = run._element.get_or_add_rPr().find(qn("w:color"))
            if color_node is None:
                color_node = OxmlElement("w:color")
                run._element.get_or_add_rPr().append(color_node)
            color_node.set(qn("w:val"), color)


def add_callout(doc, label, text, width, blue):
    label_width = 2200 if width >= 10000 else 1900
    table = doc.add_table(rows=1, cols=2)
    table.style = "Normal Table"
    set_table_geometry(table, [label_width, width - label_width])
    left, right = table.rows[0].cells
    left.text = label
    right.text = text
    set_cell_shading(left, blue)
    set_cell_shading(right, LIGHT_BLUE)
    for cell in (left, right):
        set_cell_margins(cell, top=150, bottom=150)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    style_cell(left, bold=True, color=WHITE, size=10.3)
    style_cell(right, color=TEXT, size=10.0)
    prevent_row_split(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, blue):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, blue)
        set_cell_margins(cell, top=100, bottom=100)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        style_cell(cell, bold=True, color=WHITE, size=9.6)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = value
            set_cell_margins(cell, top=95, bottom=95)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            style_cell(cell, color=TEXT, size=9.2)
        prevent_row_split(row)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code(doc, text, width):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Normal Table"
    set_table_geometry(table, [width])
    cell = table.cell(0, 0)
    cell.text = text
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=170, start=200, bottom=170, end=200)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    style_cell(cell, color="202124", size=8.4, font="DejaVu Sans Mono")
    prevent_row_split(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)
        paragraph.paragraph_format.space_after = Pt(2)


def create_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(child.get(qn("w:abstractNumId"))) for child in numbering if child.tag == qn("w:abstractNum")]
    num_ids = [int(child.get(qn("w:numId"))) for child in numbering if child.tag == qn("w:num")]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (("start", "1"), ("numFmt", "decimal"), ("lvlText", "%1."), ("suff", "tab")):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:val"), value)
        level.append(node)
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def add_steps(doc, steps):
    num_id = create_numbering(doc)
    for text in steps:
        paragraph = doc.add_paragraph(style="Normal")
        paragraph.add_run(text)
        paragraph.paragraph_format.space_after = Pt(3)
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        level = OxmlElement("w:ilvl")
        level.set(qn("w:val"), "0")
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.extend([level, num])
        p_pr.append(num_pr)


def replace_paragraph(doc, old, new):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == old:
            paragraph.text = new
            return True
    return False


def replace_paragraph_contains(doc, needle, new):
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            paragraph.text = new
            return True
    return False


def all_cells(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell


def remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_matching_paragraphs(doc, texts):
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() in texts:
            remove_paragraph(paragraph)


def remove_table_rows_by_first_cell(doc, first_cell_texts):
    for table in doc.tables:
        for row in list(table.rows):
            if row.cells and row.cells[0].text.strip() in first_cell_texts:
                table._tbl.remove(row._tr)


def remove_table_rows_by_cells_containing(doc, needles):
    for table in doc.tables:
        for row in list(table.rows):
            joined = " | ".join(cell.text.strip() for cell in row.cells)
            if any(needle in joined for needle in needles):
                table._tbl.remove(row._tr)


def append_connect_chapters(doc):
    width = 10080
    doc.add_page_break()
    doc.add_heading("27. Automatický onboarding integračním tokenem", level=1)
    doc.add_paragraph(
        "Tato kapitola je závazná pro jedinou podporovanou registraci bez ručního vyplňování v UI. "
        "Implementační token neobchází bezpečnostní brány: autorizuje jeden serverový onboardingový job, "
        "který po úplném PASS sám zaregistruje, otestuje a aktivuje právě jeden nový MCP server."
    )
    add_callout(doc, "MUST-AUTO-01", "Bez platného integračního tokenu nesmí být programátorské API dostupné. Chybějící, neplatný, expirovaný, revokovaný nebo k jinému jobu použitý token vždy vrací 401 invalid_integration_token.", width, CONNECT_BLUE)

    doc.add_heading("27.1 Vystavení a správa tokenu v UI", level=2)
    add_steps(doc, [
        "Na obrazovce Monitoring MCP nebo v sekci Implementační tokeny zvolit Vygenerovat Integrační token.",
        "Zadat srozumitelné označení. Systém vytvoří 512bitovou náhodnou hodnotu s prefixem kci_, zobrazí ji právě jednou a do databáze uloží pouze HMAC digest, key ID a šestnáctiznakový fingerprint.",
        "Token předat programátorovi odděleným bezpečným kanálem. Nevkládat jej do repozitáře, ticketu, screenshotu, logu, příkazové historie ani CI proměnné sdílené s pull requesty.",
        "V sekci Implementační tokeny sledovat aktuální expiraci, 24hodinový strop, job, KCML kód, PR/CI/deploy/testy a výslednou HTTPS adresu. Token nelze editovat; lze jej revokovat nebo skrýt s ponecháním auditu.",
    ])
    add_table(doc, ["Vlastnost", "Závazné chování"], [
        ("Entropie a formát", "64 CSPRNG bajtů = 512 bitů; kci_<Base64URL>; token se zobrazuje pouze jednou."),
        ("Počáteční TTL", "2 hodiny od vystavení. Klientské polling ani upload expiraci neprodlužují."),
        ("Serverové prodlužování", "Aktivní pronajatý job může nejvýše jednou za 15 minut posunout expiraci na now + 2h, nikdy za issuedAt + 24h."),
        ("Vazba", "Jeden token -> jeden onboarding job -> jeden KCML kód -> jeden MCP server. Druhá registrace je odmítnuta."),
        ("Po ACTIVE", "Do aktuální expirace je povolen pouze read-only GET stavu téhož jobu; nový upload, druhý server a změna identity jsou zakázány."),
        ("Resume", "Po expiraci server zůstane disabled a worker se zastaví. Správce může vystavit navazující token ke stejnému jobu; původní token se revokuje a identita se nemění."),
    ], [2600, 7480], CONNECT_BLUE)

    doc.add_heading("27.2 Rozdíl mezi revokací, zrušením a smazáním", level=2)
    add_bullets(doc, [
        "Revokovat token: programátorské API jej okamžitě odmítne; token nelze obnovit ani změnit.",
        "Zrušit job: job přejde do CANCELLED, testovací a access tokeny se revokují, server se vypne, worker a egress capability se zastaví.",
        "Smazat záznam tokenu: provede soft delete a revokaci; neměnný audit, job events a bezpečnostní důkazy zůstanou zachovány.",
        "Automatický úklid: expirované nebo revokované tokenové záznamy lze po 30 dnech skrýt; auditní události se nemažou.",
    ])

    doc.add_page_break()
    doc.add_heading("28. Programátorské API na register.hcasc.cz", level=1)
    doc.add_paragraph("Veřejné onboardingové API je oddělené od admin UI, OAuth autority a jednotlivých KCML hostů. Všechny odpovědi obsahují correlationId při chybě a stavové odpovědi používají Cache-Control: no-store.")
    doc.add_heading("28.1 První upload a rezervace identity", level=2)
    add_code(doc, """export KCML_INTEGRATION_TOKEN='kci_…'
export IDEMPOTENCY_KEY='company-project-20260713-0001'

curl --fail-with-body --request POST \\
  --url https://register.hcasc.cz/v1/onboardings \\
  --header "Authorization: Bearer ${KCML_INTEGRATION_TOKEN}" \\
  --header "Idempotency-Key: ${IDEMPOTENCY_KEY}" \\
  --form-string "manifest=$(cat onboarding-manifest.json)" \\
  --form "source=@handler.zip;type=application/zip""" , width)
    doc.add_paragraph("Úspěch vrací HTTP 202, ETag s lockVersion a tělo:")
    add_code(doc, """{
  "job": {
    "id": "<uuid>", "state": "SOURCE_UPLOADED", "lockVersion": 0,
    "code": "KCML0042", "hostname": "kcml0042.hcasc.cz",
    "resource": "https://kcml0042.hcasc.cz/mcp",
    "toolName": "kcml0042_example_handler"
  }
}""", width)
    add_callout(doc, "IDEMPOTENCE", "Opakování stejného POST se stejným tokenem, Idempotency-Key a stejnými digesty vrátí tentýž job. Stejný klíč s jiným obsahem nebo jiný klíč použitý jako druhá registrace je konflikt, nikoli nový server.", width, CONNECT_BLUE)

    doc.add_heading("28.2 Polling, opravená revize a zrušení", level=2)
    add_code(doc, """# Stav, gates, PR/CI odkazy a finální HTTPS výsledek
curl --fail-with-body \\
  --header "Authorization: Bearer ${KCML_INTEGRATION_TOKEN}" \\
  https://register.hcasc.cz/v1/onboardings/${JOB_ID}

# Opravený zdroj pouze ve stavu AWAITING_REVISION nebo FAILED
curl --fail-with-body --request PUT \\
  --url https://register.hcasc.cz/v1/onboardings/${JOB_ID}/source \\
  --header "Authorization: Bearer ${KCML_INTEGRATION_TOKEN}" \\
  --header "Idempotency-Key: ${NEW_IDEMPOTENCY_KEY}" \\
  --header 'If-Match: "<lockVersion>"' \\
  --form-string "manifest=$(cat onboarding-manifest.json)" \\
  --form "source=@handler-fixed.zip;type=application/zip"

# Zrušení téhož jobu
curl --fail-with-body --request POST \\
  --header "Authorization: Bearer ${KCML_INTEGRATION_TOKEN}" \\
  https://register.hcasc.cz/v1/onboardings/${JOB_ID}/cancel""", width)
    add_table(doc, ["HTTP", "Kód", "Význam / postup"], [
        ("400", "invalid_manifest / invalid_*", "Opravit manifest, ZIP, idempotency key nebo If-Match. Identita se nerecykluje."),
        ("401", "invalid_integration_token", "Token chybí, je chybný, expirovaný, revokovaný nebo nepatří jobu. Nezkoušet obcházet UI; vyžádat resume token."),
        ("409", "integration_token_already_bound", "Token už založil jiný obsah/job nebo operace není v aktuálním stavu povolena."),
        ("412", "lock_version_conflict", "Znovu načíst GET, převzít aktuální ETag/lockVersion a rozhodnout, zda má být odeslána nová revize."),
        ("413", "archive_too_large", "Zmenšit balík pod 10 MiB; po rozbalení musí zůstat pod 50 MiB."),
        ("202", "accepted", "Job byl přijat nebo byla idempotentně potvrzena stejná revize; pokračovat GET pollingem."),
    ], [900, 3000, 6180], CONNECT_BLUE)

    doc.add_page_break()
    doc.add_heading("29. Manifest 1.4 a zdrojový balík", level=1)
    doc.add_heading("29.1 Minimální závazná struktura", level=2)
    add_code(doc, """{
  "schemaVersion": "1.4",
  "registrationRevision": "2026-07-13.1",
  "environment": "production",
  "handlerKey": "example-handler",
  "handlerVersion": "1.0.0",
  "displayName": "Example MCP",
  "businessPurpose": "Konkrétní produkční účel nejméně o 10 znacích.",
  "owners": {"service":"team","technical":"name","security":"name","operations":"name"},
  "source": {"runtime":"nodejs22-typescript","entrypoint":"src/index.ts","testCommand":"pnpm test"},
  "runtime": {"memoryMb":128,"cpuCores":0.5,"pidsLimit":32,"egressAllowlist":["api.example.com"]},
  "tool": {
    "title":"Example", "description":"Jedna jasně vymezená funkce.",
    "inputSchema":{"type":"object","additionalProperties":false,"properties":{}},
    "outputSchema":{"type":"object","additionalProperties":false,"properties":{}},
    "annotations":{"readOnlyHint":true,"destructiveHint":false,"idempotentHint":true,"openWorldHint":false,"taskSupport":"forbidden"}
  },
  "behavior": {
    "effectClass":"READ_ONLY","timeoutMs":10000,"maxConcurrency":2,
    "requestMaxBytes":65536,"responseMaxBytes":262144,
    "rateLimit":{"windowSeconds":60,"maxRequests":30},
    "shutdownPolicy":"COMPLETE_IN_FLIGHT","idempotencyPolicy":"read-only","retryPolicy":{"automaticRetry":false}
  },
  "testContract":{"safeInput":{},"expectedResult":{},"cleanupOrCompensation":"none"},
  "monitoringProfile":{"sloTargets":{},"probeIntervals":{},"alertRules":[{"severity":"critical"}],"runbookRef":"docs/runbook.md","primaryAlertChannel":"ops","backupAlertChannel":"security"},
  "change":{"rollbackRef":"docs/rollback.md","decommissionRef":"docs/decommission.md","reviewDueAt":"2027-01-01T00:00:00.000Z"}
}""", width)
    add_callout(doc, "STRICT", "Manifest je striktní: neznámá pole, jiné runtime, automatický retry, neplatné JSON Schema nebo hodnoty mimo limity jsou odmítnuty. KCML kód, hostname, resource a finální toolName jsou výstup systému a do manifestu se nevkládají.", width, CONNECT_BLUE)

    doc.add_heading("29.2 Povinný obsah ZIP", level=2)
    add_code(doc, """handler.zip
├── package.json
├── pnpm-lock.yaml
├── tsconfig.json
└── src/
    ├── index.ts
    └── index.test.ts""", width)
    add_bullets(doc, [
        "Node.js 22, TypeScript strict, ESM. src/index.ts exportuje async function invoke(input, context).",
        "Maximálně 10 MiB komprimovaně, 50 MiB po rozbalení a 1 000 položek; cesty jsou relativní a jedinečné.",
        "Zakázány jsou symlinky, .., absolutní cesty, .env, .npmrc, .github, node_modules, Dockerfile, binární/native addony, lifecycle skripty a zakódovaná tajemství.",
        "Závislosti musí být v platformním allowlistu a mít přesné verze. git:, file:, workspace:, rozsahy a postinstall nejsou povoleny.",
        "tsconfig nesmí dědit cizí konfiguraci, používat plugins/paths/baseUrl ani zapisovat mimo dist; hlavní build ignoruje uploadované skripty.",
    ])
    doc.add_heading("29.3 Handler a řízený egress", level=2)
    add_code(doc, """type HandlerContext = Readonly<{
  correlationId: string; serverCode: string; toolName: string;
  logger: { info(fields: object, message?: string): void; error(fields: object, message?: string): void };
  egress: { fetch(url: string, init?: object): Promise<{ ok: boolean; json(): Promise<unknown> }> };
}>;

export async function invoke(input: unknown, context: HandlerContext) {
  context.logger.info({ operation: "example" }, "handler.started");
  // Jen pokud je host uveden v runtime.egressAllowlist:
  const response = await context.egress.fetch("https://api.example.com/v1", { method: "GET" });
  if (!response.ok) throw new Error("upstream_failed");
  return await response.json();
}""", width)
    doc.add_paragraph("Handler nedostává databázové, session, HMAC, GitHub, administrační, Kaja ani access-token tajemství. Nemá veřejný port a síť je deny-all. context.egress.fetch komunikuje přes privátní Unix socket s centrálním proxy, které znovu ověří capability, přesný HTTPS host/port a veřejné DNS adresy; blokuje loopback, privátní/link-local/metadata rozsahy, DNS rebinding a automatické redirecty.")

    doc.add_page_break()
    doc.add_heading("30. PR, CI, podpis a izolované nasazení", level=1)
    add_steps(doc, [
        "API ověří manifest a ZIP, provede bezpečné rozbalení, secret scan a digest a uloží balík do karantény. Webový proces uploadovaný kód nikdy nespustí.",
        "Onboarding worker s GitHub App vytvoří branch integration/kcmlNNNN/<job-id>, zapíše pouze handlers/KCMLNNNN/ a otevře PR. Workflow, root konfigurace, auth a platformní bezpečnostní soubory nesmí PR změnit. GitHub App má read oprávnění pro Actions, aby svázala úspěšný trusted main run ID s provenance.",
        "PR CI na GitHub-hosted runneru bez produkčních tajemství a bez uložených checkout credentials provede path policy, manifest/schema, lint, typecheck, unit/contract testy, secret scan, SAST, SCA/licence, SBOM a reprodukovatelný build.",
        "Po PASS se PR automaticky sloučí. Důvěryhodný main workflow použije pevný Dockerfile, ignoruje lifecycle skripty, sestaví OCI image, vytvoří SBOM a provenance, podepíše image a publikuje jej do GHCR.",
        "Worker ověří přesný source commit, build ID, image digest, podpis, SBOM a provenance. Neshoda digestu, podpisu nebo source commit je QUARANTINED bez automatického návratu.",
        "Podepsaný handler se spustí v samostatném rootless Podman workeru: non-root, read-only, cap-drop ALL, no-new-privileges, omezené CPU/RAM/PID/concurrency, log-driver none, privátní Unix socket a network none. Pevný supervisor spouští každý call v odděleném podprocesu a po registrovaném timeoutu jej ukončí, takže zablokovaný handler nemůže pokračovat na pozadí.",
    ])
    add_table(doc, ["Stav jobu", "Význam"], [
        ("CREATED / SOURCE_UPLOADED", "Token byl vystaven; poté je validovaný balík navázán na jedinou rezervovanou KCML identitu."),
        ("PR_CREATED / CI_RUNNING", "PR existuje a required checks se ukládají jako persistentní gates."),
        ("AWAITING_REVISION", "Opravitelná CI chyba; povolen je PUT nové revize se stejným tokenem/jobem a If-Match."),
        ("MERGED / ARTIFACT_BUILDING", "Zdroj je na main; čeká se na podepsaný immutable OCI artefakt."),
        ("DEPLOYING / REGISTERED_DISABLED", "Worker běží izolovaně; katalog, monitoring, statistiky, audit a digesty existují, ale veřejná funkce je disabled."),
        ("TRIAL_TESTING", "Dočasně enabled pouze pro systémové Kaja pověření a reálné veřejné testy."),
        ("ACTIVE", "Všechny gates PASS; server zůstává enabled a monitoring může vyhodnotit HEALTHY."),
        ("FAILED / QUARANTINED / CANCELLED", "Server je vypnut, access a testovací tokeny i egress capability jsou revokovány; QUARANTINED vyžaduje zásah správce."),
    ], [3000, 7080], CONNECT_BLUE)

    doc.add_heading("31. Vlastní HTTPS adresa a automatická aktivace", level=1)
    doc.add_paragraph("Každý token rezervuje vlastní identitu KCMLNNNN a vlastní resource https://kcmlNNNN.hcasc.cz/mcp. Wildcard DNS samo nestačí: produkce musí mít DNS-01 certifikát pro *.hcasc.cz, regex routing jen pro kcml[0-9]{4,}.hcasc.cz, zachovaný Host a defaultní odmítnutí jiných hostů.")
    add_table(doc, ["Gate", "Povinný PASS"], [
        ("Preflight", "DNS, TLS řetězec a SAN konkrétního hostname, SNI, Host routing, protected-resource metadata, image signature, runtime socket a worker readiness."),
        ("OAuth", "Systémové Kaja pověření jen pro testovaný server, client_credentials; chybějící a neplatný token jsou odmítnuty."),
        ("MCP", "initialize, notifications/initialized, tools/list s jediným toolName, bezpečný tools/call, input/output schema, timeout a velikostní limity."),
        ("Observabilita", "Stejné correlation ID v gateway, auditu a runtime logu; persistentní statistiky/probes; redakce a nepřítomnost integračního, Kaja, access a egress tokenu."),
        ("Monitoring", "Readiness, DNS/TLS metadata, syntetický call a artifact integrity. HEALTHY vznikne jen při úplném PASS."),
    ], [2200, 7880], CONNECT_BLUE)
    add_callout(doc, "AUTO-ACTIVE", "Po úplném PASS přepne job bez závěrečného kliknutí server na ACTIVE/enabled=true. Automatika nikdy neznamená bypass. Jakýkoli kritický bezpečnostní nebo integritní problém server vypne; digest drift, neplatný podpis nebo únik tajemství vždy znamená QUARANTINED.", width, CONNECT_BLUE)

    doc.add_heading("32. Finální checklist programátora", level=1)
    add_bullets(doc, [
        "☐ Token byl uložen jen do dočasné lokální proměnné a po skončení odstraněn; není v repozitáři, logu ani historii.",
        "☐ Manifest 1.4 je striktní, JSON Schema se kompilují a safeInput/expectedResult lze bezpečně opakovat.",
        "☐ ZIP obsahuje jen schválené kořenové soubory a TypeScript pod src/, přesný lockfile, tsconfig, src/index.ts a automatické testy.",
        "☐ První POST má stabilní Idempotency-Key; každá opravená revize má nový klíč a aktuální If-Match.",
        "☐ Polling skončil stavem ACTIVE a všechny gates jsou PASS; výsledný resource odpovídá přidělenému hostname.",
        "☐ Pokud job skončil FAILED nebo QUARANTINED, programátor stav nepřepisuje; předá correlation ID a blockingError správci.",
        "☐ Po ACTIVE byl token použit už jen pro GET stavu a poté bezpečně zničen.",
    ])
    add_callout(doc, "ACCEPT-AUTO", "Automatický onboarding je přijat pouze s jedním jobem, jedním KCML serverem, podepsaným immutable image digestem, vlastní platnou HTTPS adresou, úplnými gates PASS a finálním stavem ACTIVE/HEALTHY. Jiný výsledek není dokončená integrace.", width, CONNECT_BLUE)


def update_connect():
    doc = Document(CONNECT_SOURCE)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Verze 1.1"):
            paragraph.text = "Verze 1.2  |  13. července 2026"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif text.startswith("Tento dokument určuje přesný postup"):
            paragraph.text = (
                "Tento dokument určuje přesný postup a technické podmínky pro nový MCP server v KajovoMCPCML. "
                "Pokrývá handler, vlastní HTTPS identitu, katalog, autorizaci, logging, audit, monitoring, testy a aktivaci. "
                "Verze 1.2 zavádí jediný podporovaný automatický onboarding: správce vystaví jednorázově zobrazený integrační token, "
                "programátor odešle striktní manifest a zdrojový ZIP na register.hcasc.cz a serverový job provede PR, CI, "
                "podepsaný OCI deploy, veřejné testy a automatickou aktivaci pouze po úplném PASS."
            )
    replacements = {
        "V UI správce zvolit Přidat MCP server.": "V UI zvolit Vygenerovat Integrační token. Ruční registrace ani přímá správcovská aktivace nového serveru nejsou dostupné.",
        "Po úspěchu změnit stav přes PENDING_TECH_REVIEW, PENDING_SECURITY_REVIEW, PENDING_TEST, APPROVED, REGISTERED_DISABLED až na TRIAL nebo ACTIVE.": "Automatický job uloží jednotlivé gates, vytvoří REGISTERED_DISABLED, provede TRIAL_TESTING a po úplném PASS sám přejde do ACTIVE/enabled=true; ruční přepsání výsledku je zakázáno.",
        "☐ Server byl aktivován až po schválení a zůstává vypnutý, dokud není explicitně zapnut.": "☐ Server byl aktivován pouze automatickým stavovým automatem po úplném PASS; při automatickém onboardingu není závěrečné ruční kliknutí ani možnost obejít gate.",
        "Rezervovat identitu výhradně přes katalogové UI nebo jeho schválené API. Převzít přidělený KCML kód, hostname a resource beze změny. V handleru je nezapisovat natvrdo; handler je čte z context.server. Identita smí být uložena jen v katalogovém záznamu a schváleném manifestu.": "Použít integrační token a POST https://register.hcasc.cz/v1/onboardings. Identitu rezervuje transakčně systém; programátor převezme přidělený KCML kód, hostname, resource a toolName beze změny. Identitu nevkládá do vstupního manifestu ani natvrdo do handleru.",
        "Implementovat pouze interní KcmlHandler s key, version a invoke. Zaregistrovat jej pod přesným handlerKey@handlerVersion. Nevytvářet vlastní veřejný port, cestu, MCP server, tokenovou validaci ani alternativní vstup do funkce.": "Implementovat Node.js 22/TypeScript handler s exportovanou async function invoke a automatickými testy. Zdroj odeslat jako bezpečný ZIP; nevytvářet vlastní veřejný port, cestu, MCP server, tokenovou validaci, Dockerfile ani alternativní vstup do funkce.",
        "Vytvořit katalogový záznam ve stavu REGISTERED_DISABLED a s enabled=false. Připojit manifest, digesty, testovací vektory, výsledky CI a matici shody. Projít technickou, bezpečnostní a testovací kontrolou v předepsaném pořadí.": "Nechat onboardingový job po podpisové kontrole a izolovaném deployi vytvořit katalogový záznam REGISTERED_DISABLED/enabled=false, monitoring profile, statistiky, auditní vazby a digesty. Programátor stav ani databázi ručně nemění.",
        "Provést produkční preflight nad přidělenou HTTPS adresou: DNS, platný TLS řetězec, přesný Host routing, protected-resource metadata, autorizaci, databázi, dostupnost handleru, křížovou izolaci a fail-closed chování. Teprve po úplném PASS a explicitním schválení přejít do TRIAL nebo ACTIVE a samostatně nastavit enabled=true.": "Onboardingový job provede produkční preflight nad přidělenou HTTPS adresou, přejde do TRIAL_TESTING, vytvoří omezené systémové Kaja pověření a provede OAuth/MCP/logging/audit/monitoring testy. Po úplném PASS automaticky přejde do ACTIVE/enabled=true bez závěrečného kliknutí; jinak server vypne nebo karanténuje.",
        "Katalogový stavový automat nepovolí APPROVED, TRIAL ani ACTIVE bez platných schémat, digestů, dostupného handleru, schválení a úspěšného testovacího běhu.": "Katalogový stavový automat nepovolí TRIAL_TESTING ani ACTIVE bez platných schémat, required CI checks, podepsaného artefaktu, dostupného izolovaného workeru a úspěšného veřejného testovacího běhu.",
        "Přepínač enabled nelze aktivovat, dokud produkční preflight neověří přidělenou HTTPS adresu, TLS, routing, autorizaci a izolaci.": "enabled=true smí automaticky nastavit pouze onboardingový stavový automat po PASS DNS/TLS/routing, podpisu, autorizace, MCP kontraktu, izolace, loggingu, auditu a monitoringu.",
    }
    for old, new in replacements.items():
        replace_paragraph(doc, old, new)
    replace_paragraph_contains(doc, "apps/server/src/domain/auth.ts:", "apps/server/src/domain/auth.ts: Kaja pověření, vydání krátkodobého access tokenu a ověření oprávnění.")
    replace_paragraph_contains(doc, "Vydaný access token je krátkodobý, opaque", "Vydaný access token je krátkodobý, opaque a navázaný na schválený přístup konkrétního MCP serveru.")
    replace_paragraph_contains(doc, "☐ Access token je audience-bound", "☐ Access token je krátkodobý, opaque a vydaný pouze po splnění schválených autorizačních podmínek.")
    replace_paragraph_contains(doc, "Ověřit centrální autorizační hranici:", "Ověřit centrální autorizační hranici: resource metadata, vydání tokenu přes client_credentials, explicitní Kaja oprávnění, expiraci, revokaci, enabled stav a odmítnutí neplatného tokenu před invoke.")
    replace_paragraph_contains(doc, "Runtime vždy odvozuje resource", "Runtime vždy odvozuje přesný host z katalogu a při neshodě nebo neplatném autorizačním stavu končí fail-closed před invoke.")
    replace_paragraph_contains(doc, "Pravidelné regresní testy opakují minimálně", "Pravidelné regresní testy opakují minimálně revokaci, redakci tajemství, integritu schémat a HTTPS preflight.")
    replace_paragraph_contains(doc, "☐ Byly provedeny křížové testy mezi dvěma KCML hostname.", "☐ Byly provedeny požadované provozní a bezpečnostní testy integračního postupu.")
    doc.tables[0].cell(0, 1).text = "Závazný integrační kontrakt včetně tokenového uploadu, PR/CI, podepsaného OCI runtime, autorizace, observability, vlastní HTTPS adresy a automatické aktivace po PASS."
    remove_table_rows_by_first_cell(doc, {"T-07 Audience binding", "T-20 Dvojitá křížová izolace"})
    remove_table_rows_by_cells_containing(doc, {
        "T-07 Audience binding",
        "T-20 Dvojitá křížová izolace",
    })
    for cell in all_cells(doc):
        text = cell.text.strip()
        if text == "Audience binding již existuje":
            cell.text = "Vydávání access tokenů již existuje"
            style_cell(cell, size=9.2)
        elif text == "Access token je v databázi vázán na audience/resource a validace porovnává resource odvozený z hostname.":
            cell.text = "Access token je v databázi veden jako krátkodobý autorizační artefakt a validace kontroluje jeho stav a oprávnění."
            style_cell(cell, size=9.2)
        elif "audience-mismatched Bearer token" in text:
            cell.text = "Chybějící, neplatný nebo expirovaný Bearer token."
            style_cell(cell, size=9.2)
        elif "cizí audience" in text:
            cell.text = "Výsledky vydání tokenu, chybějícího oprávnění, expirace, revokace a disabled stavu."
            style_cell(cell, size=9.2)
        elif "cizího toolName" in text or "křížové subdomény" in text:
            cell.text = "Důkaz jednoho nástroje a odmítnutí neznámého hostu."
            style_cell(cell, size=9.2)
        if cell.text.strip().startswith("ÚKOL: Implementuj nový samostatný KCML MCP server"):
            cell.text = (
                "ÚKOL: Připrav nový samostatný KCML MCP server pro automatický onboarding do KajovoMCPCML.\n\n"
                "ZÁVAZNÝ VSTUP: Přečti celý přiložený Connect in Catalog v1.2. Vytvoř manifest 1.4 a bezpečný ZIP pro Node.js 22/TypeScript. "
                "Nevymýšlej KCML identitu; získáš ji až z POST https://register.hcasc.cz/v1/onboardings.\n\n"
                "AUTORIZACE: Použij integrační token pouze jako Bearer programátorského API. Nevkládej jej do kódu, logu, dokumentu ani commitu. "
                "První upload má stabilní Idempotency-Key; opravená revize má nový klíč a aktuální If-Match.\n\n"
                "IMPLEMENTACE: Exportuj async invoke, dodej striktní input/output schema, safeInput/expectedResult, testy a přesný lockfile. "
                "Nevytvářej veřejný port, vlastní OAuth, Dockerfile, lifecycle skript ani síťový fallback. Upstream použij jen přes context.egress.fetch a manifestový allowlist.\n\n"
                "OVĚŘENÍ: Polluj tentýž job. Za dokončené považuj pouze ACTIVE a úplné gates PASS na přidělené HTTPS adrese. "
                "FAILED, QUARANTINED, CANCELLED nebo expirovaný token jsou nedokončený výsledek; předej správci job ID, correlation ID a blockingError."
            )
            set_cell_shading(cell, LIGHT_GRAY)
            set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
            style_cell(cell, size=9.2)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip() == "1.1":
                        paragraph.text = "1.2"
    append_connect_chapters(doc)
    for paragraph in doc.paragraphs:
        if paragraph.style.name in ("Heading 1", "Heading 2", "Heading 3"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
    doc.core_properties.title = "Connect in Catalog - automatický onboarding KCML"
    doc.core_properties.subject = "Tokenový upload, PR/CI, izolovaný OCI runtime a automatická aktivace MCP serveru"
    doc.core_properties.comments = "Verze 1.2: automatický onboarding integračním tokenem, programmer API a vlastní HTTPS aktivace."
    enable_field_updates(doc)
    doc.save(CONNECT_OUTPUT)


def insert_before(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_paragraph = deepcopy(paragraph)
    paragraph._p.addprevious(new_paragraph._p)
    new_paragraph.text = text


def append_ssot_chapter(doc):
    width = 9069
    doc.add_page_break()
    doc.add_heading("37. Automatický onboarding zdrojového handleru", level=1)
    doc.add_paragraph("Tato kapitola je normativní pro jedinou podporovanou integraci jednoho nového MCP serveru integračním tokenem. Starší ruční registrační postup je odstraněn; token autorizuje workflow, nikoli přímé zapnutí serveru.")
    add_callout(doc, "AUTO-SSOT-01", "Jeden 512bitový kci_ token smí nevratně založit nebo obnovit právě jeden onboarding job a právě jednu centrálně přidělenou KCML identitu. Bez platného tokenu vrací všechny /v1/onboardings operace jednotné 401 invalid_integration_token.", width, SSOT_BLUE)
    doc.add_heading("37.1 Token, TTL a stavový automat", level=2)
    add_table(doc, ["Pravidlo", "Normativní požadavek"], [
        ("Uložení", "Jednorázové zobrazení; v DB pouze HMAC digest s odděleným klíčem, key ID, fingerprint a audit. Žádný plaintext sloupec."),
        ("TTL", "Počáteční 2h; pouze aktivní pronajatý serverový job prodlužuje v oknech na now+2h, nejvýše issuedAt+24h. Klientská aktivita neprodlužuje."),
        ("Vazba", "Token-job-server je transakční a idempotentní. Druhý server je konflikt. Po ACTIVE je povolen jen GET stavu do expirace."),
        ("Resume", "Po expiraci job zůstane blokovaný a server disabled. Nový token lze navázat na stejný job/identitu; starší tokeny se revokují."),
        ("Stavy", "CREATED, SOURCE_UPLOADED, PR_CREATED, CI_RUNNING, AWAITING_REVISION, MERGED, ARTIFACT_BUILDING, DEPLOYING, REGISTERED_DISABLED, TRIAL_TESTING, ACTIVE, FAILED, QUARANTINED, CANCELLED."),
    ], [2300, 6769], SSOT_BLUE)
    doc.add_heading("37.2 Programátorské a administrační API", level=2)
    add_bullets(doc, [
        "Admin session+CSRF: POST/GET /api/integration-tokens, GET /api/onboarding-jobs[/id], POST revoke, cancel a soft delete.",
        "Programátor: POST /v1/onboardings s Bearer tokenem, multipart manifest+source a Idempotency-Key; GET /v1/onboardings/:id; PUT /source s novým klíčem a If-Match; POST /cancel.",
        "API běží pouze na register.hcasc.cz; admin, auth, KCML a neznámý host nesmí registrační API obsloužit.",
        "POST vrací 202 a rezervuje KCMLNNNN, kcmlNNNN.hcasc.cz, přesný /mcp resource a jednoznačný toolName. Identita je výstup systému, ne vstup manifestu.",
    ])
    doc.add_heading("37.3 Intake, supply chain a runtime", level=2)
    add_bullets(doc, [
        "Manifest 1.4 je strict a popisuje Node.js 22/TypeScript ESM handler, schémata, limity, safe test, monitoring, změnové vazby a egress allowlist.",
        "ZIP má nejvýše 10 MiB/50 MiB/1000 položek a obsahuje jen schválené kořenové soubory a .ts pod src/. Traversal, symlink, binární addon, Dockerfile, .env, lifecycle script, tajemství, nepřesná/nepovolená závislost nebo rozšířený tsconfig jsou odmítnuty před spuštěním.",
        "GitHub App zapisuje jen handlers/KCMLNNNN/, vytváří PR a sleduje required checks. PR runner nemá produkční tajemství ani uložené checkout credentials. Auto-merge nastane pouze po úplném PASS. Actions-read oprávnění slouží jen ke svázání trusted main run ID s provenance.",
        "Důvěryhodný main workflow s pevným Dockerfile sestaví OCI image, SBOM a provenance, image podepíše a publikuje do GHCR. Worker ověří commit, build ID, digest, podpis a attestace.",
        "Handler běží samostatně v rootless Podman: non-root, read-only, cap-drop ALL, no-new-privileges, CPU/RAM/PID/timeout/concurrency limity, log-driver none, network none a privátní Unix socket. Pevný supervisor spouští call v podprocesu a při timeoutu jej ukončí.",
        "Strukturované logy se vracejí gateway a redigují. Povolený upstream je dostupný jen context.egress.fetch přes centrální Unix-socket proxy s per-job capability, přesným HTTPS allowlistem a SSRF/DNS-rebinding ochranou.",
    ])
    doc.add_heading("37.4 Registrace, veřejný trial a automatická aktivace", level=2)
    add_steps(doc, [
        "Po ověření OCI workeru vytvořit mcp_server jako REGISTERED_DISABLED/DISABLED/enabled=false, registration revision, monitoring profile, statistiky, audit a všechny digestové vazby.",
        "Ověřit DNS, DNS-01 wildcard certifikát *.hcasc.cz, SAN konkrétního hostu, SNI/Host routing, protected-resource metadata, podpis image a readiness Unix socketu.",
        "Přejít do TRIAL_TESTING a vytvořit krátkodobé systémové Kaja pověření s EXECUTE pouze pro testovaný server.",
        "Přes veřejné HTTPS ověřit negativní tokeny, initialize, initialized, tools/list, safe tools/call, schémata, timeout/size/rate limit, correlation, audit, logy, statistiky a probes.",
        "Revokovat systémové pověření a access tokeny. Pouze úplný PASS přepne ACTIVE/enabled=true; HEALTHY vyžaduje readiness, DNS/TLS metadata, syntetický call a artifact integrity.",
    ])
    add_callout(doc, "AUTO-SSOT-02", "Cross-host chyba, audience bypass, digest drift, neplatný podpis/provenance nebo únik integračního, Kaja, access či egress tokenu vždy nastaví QUARANTINED, revokuje tokeny/capabilities, vypne server a zastaví worker. Automatický návrat je zakázán.", width, SSOT_BLUE)
    doc.add_heading("37.5 Produkční release gate", level=2)
    add_table(doc, ["Závislost", "Release podmínka"], [
        ("HTTPS", "Wildcard DNS a platný DNS-01 certifikát *.hcasc.cz; nginx exact register host, regex KCML hosty, zachovaný Host a default deny."),
        ("GitHub", "Nainstalovaná GitHub App s minimálními contents/PR/check a Actions-read permissions; branch protection a required checks odpovídají workeru."),
        ("Supply chain", "GHCR namespace, důvěryhodný signing key, cosign verify, SBOM/provenance attestace a immutable digest."),
        ("Runtime", "Rootless Podman pro uživatele kcml, worker a egress-proxy systemd služby, karanténní/runtime adresáře a privátní socket permissions."),
        ("Testy", "Root CI včetně PostgreSQL migrací a integračních testů, onboarding PR gates, runtime/supply-chain staging E2E a desktop/mobile browser QA."),
    ], [2100, 6969], SSOT_BLUE)


def update_ssot():
    doc = Document(SSOT_SOURCE)
    remove_matching_paragraphs(doc, {
        "Automatický test cross-host izolace všech registrovaných KCML serverů.",
        "Test volání správného nástroje na správném hostu a nesprávného nástroje na správném i cizím hostu.",
    })
    remove_table_rows_by_first_cell(doc, {"ACC-HOST-01", "ACC-AUTH-02"})
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Spec Rule" and paragraph.text.startswith("Tento dokument je jediným zdrojem pravdy. Verze 1.3"):
            paragraph.text = paragraph.text.replace("Verze 1.3", "Verze 1.4", 1) + " Verze 1.4 normativně doplňuje automatický onboarding integračním tokenem, PR/CI, podepsaný izolovaný OCI runtime a automatickou aktivaci po úplném PASS."
        elif paragraph.text.strip() == "36. Konečná definice dokončení verze 1.2":
            paragraph.text = "36. Konečná definice dokončení verze 1.4"
    for cell in all_cells(doc):
        if cell.text.strip() == "1.3":
            cell.text = "1.4"
            style_cell(cell, size=9.6)
        if "Node.js 24 LTS" in cell.text:
            for paragraph in cell.paragraphs:
                if "Node.js 24 LTS" in paragraph.text:
                    paragraph.text = paragraph.text.replace("Node.js 24 LTS", "Node.js 22 LTS")
    for paragraph in doc.sections[0].header.paragraphs:
        if "SSOT v1.3" in paragraph.text:
            paragraph.text = paragraph.text.replace("SSOT v1.3", "SSOT v1.4")
    # Extend the static contents list before its trailing blank paragraph.
    toc_blank = doc.paragraphs[43]
    new_paragraph = deepcopy(doc.paragraphs[42]._p)
    toc_blank._p.addprevious(new_paragraph)
    from docx.text.paragraph import Paragraph
    Paragraph(new_paragraph, toc_blank._parent).text = "37. Automatický onboarding zdrojového handleru"
    append_ssot_chapter(doc)
    for paragraph in doc.paragraphs:
        if paragraph.style.name in ("Heading 1", "Heading 2", "Heading 3"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
    doc.core_properties.title = "KCML - Správce MCP serverů SSOT v1.4"
    doc.core_properties.subject = "Automatický onboarding, podepsaný OCI runtime a aktivace nového MCP serveru"
    doc.core_properties.comments = "Verze 1.4: tokenový onboarding, programmer API, PR/CI, OCI izolace a automatická HTTPS aktivace."
    enable_field_updates(doc)
    doc.save(SSOT_OUTPUT)


if __name__ == "__main__":
    update_connect()
    update_ssot()
    print(CONNECT_OUTPUT)
    print(SSOT_OUTPUT)

#!/usr/bin/env python3
import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[2]
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x5B, 0x65, 0x70)
LIGHT_BLUE = "E8EEF5"


def load_catalog(release: str) -> dict:
    path = ROOT / "docs" / "onboarding-catalogs" / f"component-{release}.json"
    with path.open(encoding="utf-8") as handle:
        return json.load(handle)


def output_dir(release: str) -> Path:
    directory = ROOT / "docs" / "releases" / release
    directory.mkdir(parents=True, exist_ok=True)
    return directory


def set_font(run, size=11, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = table_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for index, header in enumerate(headers):
        set_cell_fill(table.rows[0].cells[index], LIGHT_BLUE)
        paragraph = table.rows[0].cells[index].paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        set_font(paragraph.add_run(header), size=9.5, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            set_font(paragraph.add_run(str(value)), size=9.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    set_font(paragraph.add_run(text))
    return paragraph


def configure_document(doc, release):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run(f"KCML | Onboarding Catalog {release}"), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("COMPATIBLE IMPACT | katalogová změna MINOR"), size=8.5, color=MUTED)


def component_contract_rows(catalog):
    rows = []
    contracts = catalog.get("componentContracts", {})
    iterable = contracts.items() if isinstance(contracts, dict) else ((None, value) for value in contracts)
    for key, contract in iterable:
        if not isinstance(contract, dict):
            continue
        rows.append((
            contract.get("category", key or contract.get("id", "")),
            ", ".join(contract.get("requiredCapabilities", []))
            or ", ".join(contract.get("gates", []))
            or contract.get("description", ""),
        ))
    return rows[:12]


def capability_rows(catalog):
    rows = []
    contracts = catalog.get("capabilityContracts", {})
    iterable = contracts.items() if isinstance(contracts, dict) else ((None, value) for value in contracts)
    for key, capability in iterable:
        if not isinstance(capability, dict):
            continue
        rows.append((
            capability.get("capability", key or capability.get("id", "")),
            capability.get("protocol", ""),
            capability.get("requirement", capability.get("description", "required for " + ", ".join(capability.get("requiredFor", [])))),
        ))
    return rows[:12]


def compatibility_rows(catalog):
    rows = []
    for entry in catalog.get("compatibilityMatrix", []):
        rows.append((
            entry.get("profile", ""),
            entry.get("catalog", ""),
            entry.get("intake", ""),
            entry.get("result", ""),
        ))
    return rows


def runtime_rows(catalog):
    rows = []
    for section, checks in catalog.get("runtimeCompatibility", {}).items():
        for name, result in checks.items():
            rows.append((section, name, result))
    return rows


def build_docx(catalog, release, path):
    secret = catalog.get("secretManager", {})
    doc = Document()
    configure_document(doc, release)
    doc.add_paragraph().paragraph_format.space_after = Pt(90)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(kicker.add_run("KAJOVOCML - TECHNICKÝ KATALOG"), size=10, bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_font(title.add_run("Obecný komponentový model a Secret Manager"), size=27, bold=True, color=DARK_BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    set_font(subtitle.add_run(f"Onboarding Catalog {release}"), size=16, color=BLUE)
    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.space_after = Pt(80)
    set_font(lead.add_run("Kanonický model komponent, pověření, Secret API, oprávnění, auditních streamů a kompatibilních adaptérů."), size=11, italic=True, color=MUTED)
    add_table(doc, ["Vlastnost", "Hodnota"], [
        ("Dopad", "COMPATIBLE IMPACT"),
        ("Katalogová změna", "MINOR"),
        ("Release", release),
        ("MCP protokol", catalog.get("mcpProtocolVersion", "")),
        ("Manifest schema", catalog.get("manifestSchemaVersion", "")),
    ], [2700, 6660])
    doc.add_page_break()

    add_heading(doc, "1. Účel a kompatibilita")
    add_body(doc, f"Katalog {release} zachovává komponentový onboarding jako zpětně kompatibilní model a přidává centrální Secret Manager pro bezpečnou správu, rotaci a auditované zpřístupnění tajemství.")
    add_body(doc, "Stávající MCP, managed-service, Kaja a onboarding rozhraní zůstávají funkční jako adaptéry. Deaktivace komponenty technicky blokuje ingress, Pulse a egress, ale nerevokuje dlouhodobý credential; rotace a revokace jsou explicitní operace.")

    add_heading(doc, "2. Komponentový kontrakt")
    rows = component_contract_rows(catalog) or [
        ("Identita", "UUID, KCML kód a výhradní kcmlNNNN.<domain> hostname"),
        ("Autorizace", "OAuth client credentials a aktuální route/scope kontrola"),
        ("Audit", "Sekvenční stream, detekce mezer, replay, ACK a korelace"),
    ]
    add_table(doc, ["Oblast", "Kanonická pravidla"], rows, [2600, 6760])

    add_heading(doc, "3. Capability kontrakty")
    add_table(doc, ["Capability", "Protokol", "Požadavek"], capability_rows(catalog), [3300, 2300, 3760])

    add_heading(doc, "4. Secret Manager")
    public_api = secret.get("publicApi", {})
    auth = public_api.get("auth", {})
    storage = secret.get("storage", {})
    add_table(doc, ["Oblast", "Pravidlo"], [
        ("Resolve API", f"{public_api.get('hostPattern', '')} {public_api.get('resolveEndpoint', '')}".strip()),
        ("Token pro Secret API", "Bearer integration token nebo Basic client_id:client_secret; OAuth access token se pro resolve nepoužívá."),
        ("Lifecycle", "Platnost integration tokenu i client_secret je pro Secret API nezávislá na lifecycle stavu komponenty; rozhoduje autenticita credentialu a explicitní secret grant."),
        ("Reveal", "Admin reveal vyžaduje čerstvé heslo, aktuální TOTP a jednorázový grant vázaný na admina, session, secret, verzi a účel."),
        ("Úložiště", storage.get("encryption", "")),
    ], [2500, 6860])
    if auth.get("clientSecret", {}).get("tokenField") != "client_secret":
        raise ValueError("Secret Manager catalog must document client_secret as the long-lived token field.")

    doc.add_page_break()
    add_heading(doc, "5. Onboarding API v2")
    add_table(doc, ["Operace", "Endpoint", "Výsledek"], [
        ("Vytvořit", "POST /v2/component-onboardings", "Job a přidělená KCML identita"),
        ("Stav", "GET /v2/component-onboardings/{id}", "Aktuální stav a gates"),
        ("Revize", "POST /v2/component-onboardings/{id}/revisions", "Nová kanonická revize"),
        ("Readiness", "POST /v2/component-onboardings/{id}/readiness", "Vyhodnocení gates a claim token"),
        ("Credential", "POST /v2/component-onboardings/{id}/credential-claims", "Jednorázové zobrazení client_secret"),
        ("Zrušit", "DELETE /v2/component-onboardings/{id}", "Auditované zrušení"),
    ], [1900, 4400, 3060])

    add_heading(doc, "6. Autorizační důvody")
    error_rows = []
    for code in catalog.get("errorCodes", []):
        if isinstance(code, dict):
            error_rows.append((code.get("code", ""), code.get("cz", code.get("message", ""))))
        else:
            error_rows.append((str(code), "Stabilní chybový kód katalogu."))
    add_table(doc, ["Stabilní kód", "Český význam"], error_rows[:14], [3100, 6260])

    add_heading(doc, "7. Kompatibilitní matice")
    add_table(doc, ["Profil", "Katalog", "Intake / adaptér", "Výsledek"], compatibility_rows(catalog), [2450, 1700, 3300, 1910])

    add_heading(doc, "8. Runtime kompatibilita")
    add_table(doc, ["Oblast", "Kombinace", "Výsledek"], runtime_rows(catalog), [2400, 4400, 2560])

    add_heading(doc, "9. Provozní zásady")
    add_body(doc, "Gateway ověřuje Host, SNI a token audience. Interní runtime hop není zvenku adresovatelný. Autorizační rozhodnutí při každém volání čte aktuální databázový stav a aktuální route/scope oprávnění.")
    add_body(doc, "Secret Manager nikdy neukládá otevřené tajemství mimo šifrovanou verzi. Maskovaný readback, rotace, deaktivace, restore, granty, revoke a reveal jsou auditované a řízené oprávněními.")

    doc.core_properties.title = f"KajovoCML Onboarding Catalog {release}"
    doc.core_properties.subject = "Obecný komponentový model KCML a Secret Manager"
    doc.core_properties.author = "KajovoCML"
    doc.core_properties.keywords = "KCML, component, onboarding, OAuth, audit, Secret Manager"
    doc.save(path)


def build_pdf(catalog, release, path):
    styles = getSampleStyleSheet()
    title = ParagraphStyle("KcmlTitle", parent=styles["Title"], textColor=colors.HexColor("#1F4D78"), fontSize=24, leading=28)
    h2 = ParagraphStyle("KcmlH2", parent=styles["Heading2"], textColor=colors.HexColor("#2E74B5"))
    body = styles["BodyText"]
    elements = [
        Paragraph(f"KajovoCML Onboarding Catalog {release}", title),
        Paragraph("Obecný komponentový model a Secret Manager", h2),
        Paragraph("COMPATIBLE IMPACT | katalogová změna MINOR", body),
        Spacer(1, 0.2 * inch),
        Paragraph(f"Katalog {release} popisuje kanonické komponenty, kompatibilní adaptéry, Secret API, GUI správu tajemství, šifrované úložiště, rotaci, granty a auditované reveal operace.", body),
        Spacer(1, 0.15 * inch),
    ]
    rows = [["Profil", "Katalog", "Intake", "Výsledek"]] + compatibility_rows(catalog)
    table = Table(rows, repeatRows=1, colWidths=[1.55 * inch, 1.05 * inch, 2.2 * inch, 1.7 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F4D78")),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B9C3CF")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elements.extend([Paragraph("Kompatibilitní matice", h2), table, Spacer(1, 0.2 * inch)])
    secret = catalog.get("secretManager", {})
    public_api = secret.get("publicApi", {})
    storage = secret.get("storage", {})
    secret_rows = [
        ["Oblast", "Pravidlo"],
        ["Resolve API", f"{public_api.get('hostPattern', '')} {public_api.get('resolveEndpoint', '')}".strip()],
        ["Autentizace", "Bearer integration token nebo Basic client_id:client_secret; OAuth access token se nepoužívá pro resolve."],
        ["Lifecycle", "Credential autenticita a explicitní secret grant; lifecycle stav komponenty credential pro Secret API neruší."],
        ["Úložiště", storage.get("encryption", "")],
    ]
    secret_table = Table(secret_rows, repeatRows=1, colWidths=[1.55 * inch, 4.95 * inch])
    secret_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B9C3CF")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elements.extend([Paragraph("Secret Manager", h2), secret_table])
    SimpleDocTemplate(str(path), pagesize=LETTER, leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch).build(elements)


def build_markdown(catalog, release, directory):
    migration_count = "001-044" if release >= "2026.07.22" else "001-041"
    readme = f"""# KCML {release}

Katalogová změna `MINOR` s klasifikací `COMPATIBLE IMPACT`.

Release zachovává kanonický komponentový model a doplňuje centrální Secret Manager pro GUI-first správu tajemství, šifrované verze, rotaci, granty, jednorázové auditované reveal operace a Secret API. Dlouhodobý token komponenty je `client_secret`; pro Secret API je jeho platnost nezávislá na lifecycle stavu komponenty a rozhoduje autenticita credentialu společně s explicitním secret grantem.

Stávající MCP, managed-service, Kaja a onboarding adaptéry zůstávají zachované. Migrace `{migration_count}` jsou forward-only a historické katalogové artefakty se nepřepisují.

Strojově čitelné artefakty:

- `docs/onboarding-catalogs/component-{release}.json`
- `apps/server/src/contracts/component-manifest-{release}.schema.json`
- `docs/onboarding-manifest-{release}.example.json`

Lidsky čitelný katalog je v tomto adresáři ve formátech DOCX a PDF.

Úplný výsledek podporovaných legacy a nových profilů, Pulse, scope/ACL, endpoint/audience a Secret Manager kombinací je v `compatibility-matrix.md`; stejná data jsou strojově čitelná v `compatibilityMatrix`, `runtimeCompatibility` a `secretManager` katalogu.
"""
    (directory / "README.md").write_text(readme, encoding="utf-8")

    lines = [
        f"# Kompatibilitní matice KCML {release}",
        "",
        "Katalogová změna je `MINOR`, klasifikace `COMPATIBLE IMPACT`. Historické manifesty se nepřepisují; existující identity, endpointy a tokenová semantika vstupují do komponentového modelu přes adaptéry.",
        "",
        "| Profil | Katalog / manifest | Intake / adaptér | Výsledek |",
        "| --- | --- | --- | --- |",
    ]
    for profile, catalog_version, intake, result in compatibility_rows(catalog):
        lines.append(f"| `{profile}` | `{catalog_version}` | `{intake}` | `{result}` |")
    lines.extend([
        "",
        "| Runtime oblast | Kombinace | Výsledek |",
        "| --- | --- | --- |",
    ])
    for section, name, result in runtime_rows(catalog):
        lines.append(f"| `{section}` | `{name}` | `{result}` |")
    lines.extend([
        "",
        "| Secret Manager oblast | Výsledek |",
        "| --- | --- |",
        "| `integration_token` | Bearer tokeny `SINGLE_COMPONENT` a `BLUEPRINT_RELEASE` mohou volat Secret API, pokud existuje explicitní grant. |",
        "| `client_secret` | Dlouhodobý token `client_secret` je ověřen přímo přes Basic `client_id:client_secret` a pro Secret API není nahrazen OAuth access tokenem. |",
        "| `component_lifecycle` | `DISABLED`, `INACTIVE`, `QUARANTINED` a `DEREGISTERED` samy o sobě credential pro Secret API neruší; chybějící grant nebo credential selže fail-closed. |",
        "| `admin_reveal` | Vyžaduje čerstvé heslo, aktuální TOTP a jednorázový grant vázaný na admina, session, secret, verzi a účel. |",
    ])
    (directory / "compatibility-matrix.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def build(release: str):
    catalog = load_catalog(release)
    if catalog.get("catalogVersion") != release:
        raise ValueError(f"Catalog version mismatch: expected {release}, got {catalog.get('catalogVersion')}")
    if release >= "2026.07.22" and not catalog.get("secretManager"):
        raise ValueError("Secret Manager section is required for 2026.07.22+ release documents.")
    directory = output_dir(release)
    build_docx(catalog, release, directory / f"KajovoCML_Onboarding_Catalog_{release}.docx")
    build_pdf(catalog, release, directory / f"KajovoCML_Onboarding_Catalog_{release}.pdf")
    build_markdown(catalog, release, directory)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Build KCML human-readable onboarding catalog documents.")
    parser.add_argument("--release", default="2026.07.21")
    args = parser.parse_args()
    build(args.release)
